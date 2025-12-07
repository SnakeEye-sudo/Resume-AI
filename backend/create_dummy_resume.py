from docx import Document

def create_dummy_resume(filename="dummy_resume.docx"):
    doc = Document()
    doc.add_heading('John Doe', 0)
    
    doc.add_paragraph('Email: john.doe@example.com | Phone: (555) 123-4567')
    doc.add_paragraph('Location: New York, NY')
    
    doc.add_heading('Skills', level=1)
    doc.add_paragraph('Python, FastAPI, React, JavaScript, SQL, Docker, AWS, Machine Learning')
    
    doc.add_heading('Experience', level=1)
    doc.add_heading('Senior Software Engineer - Tech Corp', level=2)
    doc.add_paragraph('Jan 2020 - Present')
    doc.add_paragraph('Developed scalable backend APIs using FastAPI and PostgreSQL.')
    doc.add_paragraph('Implemented CI/CD pipelines with Docker and GitHub Actions.')
    
    doc.add_heading('Education', level=1)
    doc.add_paragraph('B.S. Computer Science - University of Tech (2019)')
    
    doc.save(filename)
    print(f"Created {filename}")

if __name__ == "__main__":
    create_dummy_resume()
